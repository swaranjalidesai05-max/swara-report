from flask import (
    Flask,
    render_template,
    request,
    redirect,
    url_for,
    session,
    flash,
    send_file,
    jsonify,
)
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from datetime import datetime
import sqlite3
import os
import json
import csv
import io
import requests
import matplotlib
matplotlib.use('Agg')  # Non-interactive backend for server-side rendering
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


app = Flask(__name__)
app.secret_key = "change-this-secret"

UPLOAD_FOLDER = "static"
# Use the actual template present in the project
TEMPLATE_PATH = "word_templates/college_letterhead.docx"
GENERATED_FOLDER = "generated_reports"

# ---------------- PSO CONSTANTS ----------------
# Programme Specific Outcomes text used in generated reports.
PSO1_TEXT = (
    "PSO1: An ability to apply the theoretical concepts and practical knowledge of "
    "Information Technology in the analysis, design, development, and management of "
    "information processing systems and applications in the interdisciplinary domain "
    "to understand professional, business processes, ethical, legal, security, and "
    "social issues and responsibilities."
)

PSO2_TEXT = (
    "PSO2: An ability to analyze a problem and identify and define the computing "
    "infrastructure and operations requirements appropriate to its solution. IT "
    "graduates should be able to work on large-scale computing systems."
)

# Programme Outcomes (PO) – hardcoded headings only; no custom PO.
# Used in multi-select "selected_pos" and in report generation.
PO_HEADINGS = [
    "Engineering Knowledge",
    "Problem Analysis",
    "Design / Development of Solutions",
    "Conduct investigations of complex problems",
    "Modern Tool Usage",
    "The Engineer and Society",
    "Environment and Sustainability",
    "Ethics",
    "Communication",
    "Project Management & Finance",
    "Lifelong Learning",
]

# Sustainable Development Goals (SDGs) – predefined list for checklist and report.
# List of (code, title) for 17 SDGs. Stored as selected_sdgs JSON array of codes.
SDG_LIST = [
    ("SDG1", "No Poverty"),
    ("SDG2", "Zero Hunger"),
    ("SDG3", "Good Health and Well-being"),
    ("SDG4", "Quality Education"),
    ("SDG5", "Gender Equality"),
    ("SDG6", "Clean Water and Sanitation"),
    ("SDG7", "Affordable and Clean Energy"),
    ("SDG8", "Decent Work and Economic Growth"),
    ("SDG9", "Industry, Innovation and Infrastructure"),
    ("SDG10", "Reduced Inequalities"),
    ("SDG11", "Sustainable Cities and Communities"),
    ("SDG12", "Responsible Consumption and Production"),
    ("SDG13", "Climate Action"),
    ("SDG14", "Life Below Water"),
    ("SDG15", "Life on Land"),
    ("SDG16", "Peace, Justice and Strong Institutions"),
    ("SDG17", "Partnerships for the Goals"),
]


# ---------------- DATABASE ----------------

def init_db():
    conn = sqlite3.connect("database.db")
    c = conn.cursor()

    # Users table (with email and role)
    c.execute(
        """
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE,
            email TEXT,
            password TEXT
        )
        """
    )

    # Ensure email / role columns exist on old databases
    c.execute("PRAGMA table_info(users)")
    user_cols = [row[1] for row in c.fetchall()]
    if "email" not in user_cols:
        c.execute("ALTER TABLE users ADD COLUMN email TEXT")
    if "role" not in user_cols:
        # Default everyone to 'Student' unless explicitly changed later
        c.execute("ALTER TABLE users ADD COLUMN role TEXT DEFAULT 'Student'")
        # Backfill any existing rows to have a role
        c.execute("UPDATE users SET role='Student' WHERE role IS NULL OR role=''")

    # Events table – stores all event metadata
    c.execute(
        """
        CREATE TABLE IF NOT EXISTS events (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT,
            date TEXT,
            venue TEXT,
            department TEXT,
            description TEXT,
            event_photo TEXT,
            academic_year TEXT,
            resource_person TEXT,
            resource_designation TEXT,
            event_coordinator TEXT,
            event_time TEXT,
            event_type TEXT,
            permission_letter TEXT,
            invitation_letter TEXT,
            notice_letter TEXT,
            appreciation_letter TEXT,
            event_photos TEXT,
            attendance_photo TEXT,
            outcome_1 TEXT,
            outcome_2 TEXT,
            outcome_3 TEXT,
            feedback_data TEXT,
            pso1_selected INTEGER DEFAULT 0,
            pso2_selected INTEGER DEFAULT 0,
            selected_pos TEXT,
            selected_sdgs TEXT
        )
        """
    )

    # Add any missing columns for existing databases
    c.execute("PRAGMA table_info(events)")
    event_cols = [row[1] for row in c.fetchall()]
    new_event_columns = [
        ("event_photo", "TEXT"),
        ("academic_year", "TEXT"),
        ("resource_person", "TEXT"),
        ("resource_designation", "TEXT"),
        ("event_coordinator", "TEXT"),
        ("event_time", "TEXT"),
        ("event_type", "TEXT"),
        ("permission_letter", "TEXT"),
        ("invitation_letter", "TEXT"),
        ("notice_letter", "TEXT"),
        ("appreciation_letter", "TEXT"),
        ("event_photos", "TEXT"),
        ("attendance_photo", "TEXT"),
        ("outcome_1", "TEXT"),
        ("outcome_2", "TEXT"),
        ("outcome_3", "TEXT"),
        ("feedback_data", "TEXT"),
        ("pso1_selected", "INTEGER DEFAULT 0"),
        ("pso2_selected", "INTEGER DEFAULT 0"),
        ("selected_pos", "TEXT"),
        ("selected_sdgs", "TEXT"),
        ("feedback_form_link", "TEXT"),
        ("feedback_sheet_id", "TEXT"),
    ]
    for col_name, col_type in new_event_columns:
        if col_name not in event_cols:
            c.execute(f"ALTER TABLE events ADD COLUMN {col_name} {col_type}")

    # Reports table – stores generated report info
    c.execute(
        """
        CREATE TABLE IF NOT EXISTS reports (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            event_id INTEGER,
            file_path TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        """
    )

    # Ensure newer analytics-related columns exist on old databases
    c.execute("PRAGMA table_info(reports)")
    report_cols = [row[1] for row in c.fetchall()]
    if "status" not in report_cols:
        # Basic lifecycle tracking for reports; default everything to 'submitted'
        c.execute("ALTER TABLE reports ADD COLUMN status TEXT DEFAULT 'submitted'")
        c.execute("UPDATE reports SET status='submitted' WHERE status IS NULL OR status=''")

    conn.commit()
    conn.close()


def get_db():
    conn = sqlite3.connect("database.db")
    conn.row_factory = sqlite3.Row
    return conn


# --------------- FILE HELPERS ---------------

def _save_file(file_storage, subfolder, allow_pdf=False):
    """Save an uploaded file and return its relative path, or None."""
    if not file_storage or file_storage.filename == "":
        return None

    filename = secure_filename(file_storage.filename)
    ext = os.path.splitext(filename)[1].lower()
    image_exts = {".jpg", ".jpeg", ".png", ".gif"}
    allowed_exts = image_exts | ({".pdf"} if allow_pdf else set())
    if ext not in allowed_exts:
        return None

    os.makedirs(os.path.join(UPLOAD_FOLDER, subfolder), exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
    filename = f"{timestamp}_{filename}"
    rel_path = os.path.join(UPLOAD_FOLDER, subfolder, filename)
    abs_path = os.path.join(rel_path)
    file_storage.save(abs_path)
    return rel_path.replace("\\", "/")


# ---------------- AUTH ----------------

def login_required(f):
    from functools import wraps
    @wraps(f)
    def wrap(*args, **kwargs):
        if "user_id" not in session:
            return redirect(url_for("login"))
        return f(*args, **kwargs)
    return wrap


def hod_required_api(f):
    """
    Simple decorator for HOD-only JSON APIs.
    Assumes user is already authenticated via @login_required.
    """
    from functools import wraps

    @wraps(f)
    def wrap(*args, **kwargs):
        if session.get("role") != "HOD":
            # Return a JSON 403 so frontend can handle gracefully
            return jsonify({"error": "Forbidden", "message": "HOD access required"}), 403
        return f(*args, **kwargs)

    return wrap


# ---------------- DOCX HELPERS ----------------
from docx.shared import Pt, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import copy

# --------------- Shared font helper ---------------

def _apply_font(run, name="Times New Roman", size=Pt(12), bold=False):
    """Apply consistent font properties to a run."""
    run.font.name = name
    run.font.size = size
    run.bold = bold
    # Force font in East-Asian contexts as well
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def _find_marker_paragraphs(doc, marker):
    """Yield all paragraphs (body, tables, headers/footers) containing *marker*."""
    for p in doc.paragraphs:
        if marker in "".join(r.text for r in p.runs) or marker in p.text:
            yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if marker in "".join(r.text for r in p.runs) or marker in p.text:
                        yield p
    for section in doc.sections:
        for hf in (section.header, section.footer):
            for p in hf.paragraphs:
                if marker in "".join(r.text for r in p.runs) or marker in p.text:
                    yield p
            for table in hf.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if marker in "".join(r.text for r in p.runs) or marker in p.text:
                                yield p


def create_default_template():
    """Create a default Word template with PO, PSO, and SDG placeholders if missing."""
    if os.path.exists(TEMPLATE_PATH):
        return
    os.makedirs(os.path.dirname(TEMPLATE_PATH), exist_ok=True)
    doc = Document()
    doc.add_paragraph("Programme Specific Outcomes (PSO)")
    doc.add_paragraph("{{PSO_SECTION}}")
    doc.add_paragraph()
    doc.add_paragraph("Program Outcomes (PO)")
    doc.add_paragraph("{{PO_SECTION}}")
    doc.add_paragraph()
    doc.add_paragraph("Sustainable Development Goals (SDGs) Addressed")
    doc.add_paragraph("{{SDG_SECTION}}")
    doc.save(TEMPLATE_PATH)


# --------------- 1. Placeholder replacement (format-preserving) ---------------

def replace_placeholders(doc, replacements):
    """
    Replace {{placeholder}} markers while preserving run-level formatting.

    Strategy:
      - Reconstruct the full paragraph text from all runs.
      - If a placeholder spans multiple runs (Word sometimes splits text),
        we locate which runs participate, coalese them, perform the
        replacement, and keep the *first* participating run's formatting
        for the result.
      - Single-run replacements simply swap the text inside the run,
        keeping font, bold, colour, size intact.
    """
    def _replace_in_paragraph(para):
        for placeholder, value in replacements.items():
            value = str(value or "")
            # Keep trying until every occurrence in this paragraph is handled
            safety = 20  # prevent infinite loop
            while safety > 0:
                safety -= 1
                full_text = "".join(run.text for run in para.runs)
                start = full_text.find(placeholder)
                if start == -1:
                    break

                end = start + len(placeholder)

                # Map character positions to runs
                char_idx = 0
                first_run_idx = last_run_idx = None
                for ri, run in enumerate(para.runs):
                    run_start = char_idx
                    run_end = char_idx + len(run.text)
                    if first_run_idx is None and run_end > start:
                        first_run_idx = ri
                    if run_end >= end and last_run_idx is None:
                        last_run_idx = ri
                        break
                    char_idx = run_end

                if first_run_idx is None or last_run_idx is None:
                    break

                # Build prefix (text before placeholder in first run)
                first_run = para.runs[first_run_idx]
                offset_in_first = start - sum(len(para.runs[i].text) for i in range(first_run_idx))
                prefix = first_run.text[:offset_in_first]

                # Build suffix (text after placeholder in last run)
                last_run = para.runs[last_run_idx]
                consumed_before_last = sum(len(para.runs[i].text) for i in range(last_run_idx))
                offset_in_last = end - consumed_before_last
                suffix = last_run.text[offset_in_last:]

                # Set the first run text to prefix + replacement value + suffix
                first_run.text = prefix + value + suffix

                # Remove intermediate + last runs that were part of placeholder
                for ri in range(last_run_idx, first_run_idx, -1):
                    para._p.remove(para.runs[ri]._r)

    # Body paragraphs
    for p in doc.paragraphs:
        _replace_in_paragraph(p)

    # Table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p)

    # Headers & footers
    for section in doc.sections:
        for hf in (section.header, section.footer):
            for p in hf.paragraphs:
                _replace_in_paragraph(p)
            for table in hf.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            _replace_in_paragraph(p)


# --------------- 2. Event details as a clean table ---------------

def _add_detail_run(para, text, bold=False, italic=False):
    """Add a run with Times New Roman 12pt to a paragraph."""
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    # Force east-asian font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return run


def insert_event_details_paragraph(doc, marker, event):
    """
    Replace the marker paragraph with event details formatted as clean
    paragraphs matching the college report style:
      - Line 1: "Academic Year: XXXX"  (tab)  "Date: XXXX"
      - Line 2: "Name of the Event: XXXX"
      - Line 3: "Resource Person: XXXX"
      - etc.
    Labels are bold+italic, values are regular.
    """
    # Fields displayed one per line (label: value)
    detail_lines = [
        ("Name of the Event: ", event.get("title", "")),
        ("Resource Person: ", event.get("resource_person", "")),
        ("Designation: ", event.get("resource_designation", "")),
        ("Event Type: ", event.get("event_type", "")),
        ("Time: ", event.get("event_time", "")),
        ("Venue: ", event.get("venue", "")),
        ("Department: ", event.get("department", "")),
        ("Event Coordinator: ", event.get("event_coordinator", "")),
    ]

    for p in doc.paragraphs:
        if marker not in p.text:
            continue

        anchor = p._p

        # We'll build paragraphs in reverse order and insert after anchor,
        # so the final order in the document is correct.
        paras_to_insert = []

        # --- First line: Academic Year + Date (same line, tab-separated) ---
        first_para = doc.add_paragraph()
        first_para.paragraph_format.space_before = Pt(6)
        first_para.paragraph_format.space_after = Pt(2)
        first_para.paragraph_format.line_spacing = 1.15
        _add_detail_run(first_para, "Academic Year: ", bold=True, italic=False)
        _add_detail_run(first_para, str(event.get("academic_year", "")))
        _add_detail_run(first_para, "\t\t\t\t\t")
        _add_detail_run(first_para, "Date:  ", bold=True, italic=False)
        _add_detail_run(first_para, str(event.get("date", "")))
        paras_to_insert.append(first_para)

        # --- Remaining fields: one per line ---
        for label, value in detail_lines:
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.line_spacing = 1.15
            _add_detail_run(para, label, bold=True, italic=False)
            _add_detail_run(para, str(value or ""))
            paras_to_insert.append(para)

        # Move all created paragraphs from end of doc to after the anchor
        # Insert in reverse so order is preserved
        for para in reversed(paras_to_insert):
            anchor.addnext(para._p)

        # Remove the original marker paragraph
        anchor.getparent().remove(anchor)
        return


# --------------- 3. Full-page image insertion (scanned documents) ---------------

def insert_full_page_image(doc, marker, path, heading="", page_break=True):
    """
    Replace marker paragraph with a section containing:
      1. Page break (only if page_break=True AND there is real content before)
      2. Bold centered heading (16pt, Times New Roman)
      3. Centered image below the heading

    If path is missing or file doesn't exist, the marker paragraph
    is silently removed so no placeholder text remains.

    Set page_break=False to group multiple images on the same page flow
    (no page break inserted before this image).
    """
    for p_idx, p in enumerate(doc.paragraphs):
        full_text = "".join(r.text for r in p.runs) or p.text
        if marker not in full_text:
            continue

        # If no file, remove marker paragraph cleanly
        if not path or not os.path.exists(path):
            p._p.getparent().remove(p._p)
            return

        # Check if there is meaningful content before this paragraph.
        # If the marker is at/near the top (no real content before it),
        # skip the page break to avoid an empty first page.
        has_content_before = False
        for prev_p in doc.paragraphs[:p_idx]:
            prev_text = prev_p.text.strip()
            # Ignore empty paragraphs and other unprocessed markers
            if prev_text and not prev_text.startswith("<<") and not prev_text.startswith("{{"):
                has_content_before = True
                break

        # Clear marker text
        p.clear()

        # --- Page break (only if enabled AND there's content before) ---
        if page_break and has_content_before:
            break_run = p.add_run()
            break_run.add_break(WD_BREAK.PAGE)

        # --- Section heading ---
        if heading:
            title_run = p.add_run(heading)
            title_run.bold = True
            title_run.font.name = "Times New Roman"
            title_run.font.size = Pt(16)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)

        # --- Image in its own paragraph below the heading ---
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_para.paragraph_format.keep_together = True
        img_para.paragraph_format.space_before = Pt(6)
        img_para.paragraph_format.space_after = Pt(6)
        run = img_para.add_run()
        run.add_picture(path, width=Inches(5.8))
        # Move image paragraph from end of doc to right after the heading
        p._p.addnext(img_para._p)
        return

    # Also check tables (less common)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if marker in p.text:
                        if not path or not os.path.exists(path):
                            p._p.getparent().remove(p._p)
                            return
                        p.clear()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        run.add_picture(path, width=Inches(5.5))
                        return


# --------------- 4. Event photos ---------------

def insert_event_photos(doc, marker, photos):
    """
    Replace marker with event photos placed inline.
    The template already has a "Photographs:" heading before the marker,
    so we only insert the images — no duplicate heading, no page break.
    Each image: centered, consistent width (5 inches), spacing between.
    """
    if not photos:
        # Remove the marker paragraph entirely
        for p in doc.paragraphs:
            if marker in p.text:
                p._p.getparent().remove(p._p)
                return
        return

    # Filter to only existing files
    valid_photos = [img for img in photos if os.path.exists(img)]
    if not valid_photos:
        for p in doc.paragraphs:
            if marker in p.text:
                p._p.getparent().remove(p._p)
                return
        return

    for p in doc.paragraphs:
        if marker not in p.text:
            continue

        # Use marker paragraph as anchor, then remove it
        anchor = p._p

        # Insert each photo in its own paragraph (all in same section)
        prev_element = anchor
        for img_path in valid_photos:
            img_para = doc.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_para.paragraph_format.keep_together = True
            img_para.paragraph_format.space_before = Pt(6)
            img_para.paragraph_format.space_after = Pt(18)
            run = img_para.add_run()
            run.add_picture(img_path, width=Inches(5))
            # Move from end of doc to right after the previous element
            prev_element.addnext(img_para._p)
            prev_element = img_para._p

        # Remove original marker paragraph
        anchor.getparent().remove(anchor)
        return


# --------------- 5. Attendance ---------------

def insert_attendance(doc, marker, path):
    """
    Replace marker with a NEW PAGE containing:
      1. Page break
      2. "Attendance" heading (bold, centered, 16pt)
      3. Attendance image or note
    """
    if not path or not os.path.exists(path):
        # Remove marker paragraph if no attendance file
        for p in doc.paragraphs:
            if marker in p.text:
                p._p.getparent().remove(p._p)
                return
        return

    for p in doc.paragraphs:
        if marker not in p.text:
            continue

        p.clear()

        # --- Page break ---
        break_run = p.add_run()
        break_run.add_break(WD_BREAK.PAGE)

        # --- "Attendance" heading ---
        title_run = p.add_run("Attendance")
        title_run.bold = True
        title_run.font.name = "Times New Roman"
        title_run.font.size = Pt(16)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)

        ext = os.path.splitext(path)[1].lower()

        if ext in (".jpg", ".jpeg", ".png", ".gif"):
            # Image in its own paragraph below heading
            img_para = doc.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_para.paragraph_format.keep_together = True
            img_para.paragraph_format.space_before = Pt(6)
            run = img_para.add_run()
            run.add_picture(path, width=Inches(6.0))
            # Move paragraph to right after the heading
            p._p.addnext(img_para._p)
        else:
            p.add_run("\n")
            note_run = p.add_run("Attendance attached as scanned document.")
            _apply_font(note_run, size=Pt(12))

        return


# --------------- 6. Feedback Form & Analysis ---------------

import re

def _extract_sheet_id(raw_value):
    """
    Extract the Google Sheet ID from a full URL or return the value as-is
    if it already looks like a bare ID.

    Accepts any of:
      - https://docs.google.com/spreadsheets/d/SHEET_ID/edit#gid=0
      - https://docs.google.com/spreadsheets/d/SHEET_ID/
      - https://docs.google.com/spreadsheets/d/SHEET_ID
      - SHEET_ID  (bare, already extracted)
    """
    if not raw_value:
        return ""
    raw_value = raw_value.strip()
    # Try to extract from a URL
    m = re.search(r"/spreadsheets/d/([a-zA-Z0-9_-]+)", raw_value)
    if m:
        extracted = m.group(1)
        print(f"[Feedback] Extracted Sheet ID from URL: {extracted}")
        return extracted
    # Already a bare ID (no slashes, looks like an alphanumeric string)
    return raw_value


def fetch_google_sheet_data(sheet_id):
    """
    Fetch feedback responses from a Google Sheet linked to a Google Form.

    Accepts a bare Sheet ID or a full Google Sheets URL (auto-extracted).

    Tries two methods (no API key required):
      1. Direct CSV export  — works if sheet is "Published to Web"
      2. Google Visualization API (gviz/tq) — works if sheet is shared via link

    Returns a list of dicts (one per row), or an empty list on failure.
    """
    # Auto-extract the ID if the user pasted a full URL
    sheet_id = _extract_sheet_id(sheet_id)
    if not sheet_id:
        print("[Feedback] No Sheet ID provided")
        return []

    print(f"[Feedback] Fetching data for Sheet ID: {sheet_id}")

    urls = [
        # Method 1: direct export (published sheets)
        f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=csv",
        # Method 2: gviz endpoint (link-shared sheets)
        f"https://docs.google.com/spreadsheets/d/{sheet_id}/gviz/tq?tqx=out:csv",
    ]
    for url in urls:
        try:
            resp = requests.get(url, timeout=20)
            resp.raise_for_status()
            resp.encoding = "utf-8"
            text = resp.text.strip()
            if not text or text.startswith("<!"):  # HTML error page
                print(f"[Feedback] Got HTML/empty response from {url[:60]}…")
                continue
            reader = csv.DictReader(io.StringIO(text))
            rows = [row for row in reader]
            if rows:
                print(f"[Feedback] Fetched {len(rows)} responses from Google Sheet")
                return rows
        except Exception as e:
            print(f"[Feedback] Endpoint failed ({url[:60]}…): {e}")
            continue
    print("[Feedback] Could not fetch data from any endpoint")
    return []


def _find_column(rows, *keywords):
    """
    Find a column header that matches any of the keywords (case-insensitive
    partial match). Returns the exact header string or None.
    """
    if not rows:
        return None
    for hdr in rows[0].keys():
        hdr_lower = hdr.lower().strip()
        for kw in keywords:
            if kw in hdr_lower:
                return hdr
    return None


def _count_column(rows, col):
    """Count non-empty values in *col* across all *rows*. Returns an OrderedDict."""
    counts = {}
    for r in rows:
        val = (r.get(col) or "").strip()
        if val:
            counts[val] = counts.get(val, 0) + 1
    return counts


def generate_feedback_charts(responses, output_dir):
    """
    Generate high-resolution Pie Charts for EVERY categorical column
    detected in Google Sheet feedback responses.

    Auto-detection logic:
      1. Skips columns whose header contains keywords like 'timestamp',
         'email', 'name' (these are identifiers, not categorical data).
      2. Skips columns where > 60% of values are unique (likely free-text).
      3. Skips columns with fewer than 2 non-empty responses.
      4. Skips columns with more than 15 distinct values (not categorical).
      5. Everything else gets a Pie Chart with percentage labels & legend.

    Returns:
        (chart_info_list, stats)
        chart_info_list — list of dicts: {"path": str, "title": str}
        stats           — dict with summary statistics for the Word report
    """
    os.makedirs(output_dir, exist_ok=True)
    chart_info = []  # {"path": ..., "title": ...}
    stats = {"total": len(responses), "charts_generated": 0}

    if not responses:
        return chart_info, stats

    total = len(responses)

    # ---- Professional colour palette (Tableau-inspired) ----
    COLOR_PALETTE = [
        "#4e79a7", "#f28e2b", "#e15759", "#76b7b2", "#59a14f",
        "#edc948", "#b07aa1", "#ff9da7", "#9c755f", "#bab0ac",
        "#6b6ecf", "#d6616b", "#8ca252", "#de9ed6", "#ad494a",
    ]

    # ---- Column headers to skip (identifiers / non-categorical) ----
    SKIP_KEYWORDS = [
        "timestamp", "email", "e-mail", "mail id",
        "name", "full name", "phone", "mobile",
        "roll", "prn", "enroll",
        "address", "comment", "suggestion", "feedback",
        "remark", "describe", "explain", "write",
        "other", "specify",
    ]

    def _should_skip_header(header):
        """Return True if the column header matches a skip keyword."""
        h = header.lower().strip()
        for kw in SKIP_KEYWORDS:
            if kw in h:
                return True
        return False

    def _is_categorical(rows, col):
        """
        Determine if a column is categorical (suitable for a pie chart).
        Returns (True, counts_dict) or (False, None).
        """
        counts = {}
        non_empty = 0
        for r in rows:
            val = (r.get(col) or "").strip()
            if val:
                non_empty += 1
                counts[val] = counts.get(val, 0) + 1
        if non_empty < 2:
            return False, None
        distinct = len(counts)
        if distinct < 2 or distinct > 15:
            return False, None
        # If more than 60% of values are unique → free-text, skip
        if distinct / non_empty > 0.6:
            return False, None
        return True, counts

    # ---- Iterate through ALL columns and generate Pie Charts ----
    all_columns = list(responses[0].keys()) if responses else []
    chart_index = 0

    for col in all_columns:
        # Skip non-categorical headers
        if _should_skip_header(col):
            continue

        # Check if column data is categorical
        is_cat, counts = _is_categorical(responses, col)
        if not is_cat or not counts:
            continue

        # --- Build the Pie Chart ---
        labels = list(counts.keys())
        sizes = list(counts.values())
        total_for_col = sum(sizes)
        colors = [COLOR_PALETTE[i % len(COLOR_PALETTE)] for i in range(len(labels))]

        # Build legend labels: "Option — count (percentage%)"
        legend_labels = [
            f"{lbl} - {cnt} ({cnt / total_for_col * 100:.1f}%)"
            for lbl, cnt in zip(labels, sizes)
        ]

        # Create figure — 600×400 equivalent at 100 DPI = figsize (8, 5) at 200 DPI
        fig, ax = plt.subplots(figsize=(8, 5), dpi=200)

        wedges, texts, autotexts = ax.pie(
            sizes,
            autopct="%1.1f%%",
            colors=colors,
            startangle=140,
            pctdistance=0.78,
            wedgeprops={"edgecolor": "white", "linewidth": 1.5},
            textprops={"fontsize": 9, "fontfamily": "serif"},
        )
        for at in autotexts:
            at.set_fontsize(9)
            at.set_fontweight("bold")
            at.set_color("white")

        # Title = column header (the Google Form question)
        chart_title = col.strip()
        ax.set_title(chart_title, fontsize=14,
                     fontfamily="serif", fontweight="bold", pad=16)

        # Legend on the right side
        ax.legend(
            wedges, legend_labels,
            title="Responses",
            loc="center left",
            bbox_to_anchor=(1.0, 0.5),
            fontsize=9,
            title_fontsize=10,
            frameon=False,
        )

        # Subtitle with total responses for this question
        ax.text(0, -1.25, f"Total Responses: {total_for_col}",
                ha="center", fontsize=9, fontfamily="serif", color="#555555")

        fig.tight_layout()

        # Save as PNG
        safe_name = "".join(c if c.isalnum() or c == "_" else "_" for c in col.lower())
        filename = f"chart_{chart_index:02d}_{safe_name[:40]}.png"
        path = os.path.join(output_dir, filename)
        fig.savefig(path, bbox_inches="tight", facecolor="white")
        plt.close(fig)

        chart_info.append({"path": path, "title": chart_title})
        chart_index += 1

    stats["charts_generated"] = len(chart_info)
    stats["columns_detected"] = [ci["title"] for ci in chart_info]
    print(f"[Feedback] Generated {len(chart_info)} pie charts from {len(all_columns)} columns")
    return chart_info, stats


def insert_feedback_analysis(doc, marker, feedback_form_link, feedback_sheet_id, charts_dir):
    """
    Replace the <<FEEDBACK_TABLE>> marker with a 'Feedback Form & Analysis' section:
      1. Clickable hyperlink to the Google Form
      2. Summary statistics paragraph (total responses + detected columns)
      3. For EACH categorical column: a sub-heading + its Pie Chart image

    If no form link / sheet ID is provided, the marker paragraph is removed silently.
    Charts are generated fresh every time so they always reflect the latest data.
    No hardcoded column names — everything is dynamic.
    """
    has_data = feedback_form_link or feedback_sheet_id
    if not has_data:
        for p in doc.paragraphs:
            if marker in p.text:
                p._p.getparent().remove(p._p)
                return
        return

    # Fetch live data and build charts
    chart_info = []
    stats = {"total": 0}
    fetch_failed = False
    if feedback_sheet_id:
        responses = fetch_google_sheet_data(feedback_sheet_id)
        if responses:
            chart_info, stats = generate_feedback_charts(responses, charts_dir)
        else:
            fetch_failed = True  # Sheet exists but data couldn't be retrieved

    for p in doc.paragraphs:
        if marker not in p.text:
            continue

        anchor = p._p
        paras_to_insert = []

        # --- 1. Clickable hyperlink to Google Form ---
        if feedback_form_link:
            link_para = doc.add_paragraph()
            link_para.paragraph_format.space_before = Pt(12)
            link_para.paragraph_format.space_after = Pt(6)

            label_run = link_para.add_run("Feedback Form: ")
            _apply_font(label_run, bold=True, size=Pt(12))

            # Build clickable hyperlink
            r_id = doc.part.relate_to(
                feedback_form_link,
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                is_external=True,
            )
            hyperlink = OxmlElement("w:hyperlink")
            hyperlink.set(qn("r:id"), r_id)

            new_run = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")

            u_elem = OxmlElement("w:u")
            u_elem.set(qn("w:val"), "single")
            rPr.append(u_elem)

            clr = OxmlElement("w:color")
            clr.set(qn("w:val"), "0563C1")
            rPr.append(clr)

            rFonts = OxmlElement("w:rFonts")
            rFonts.set(qn("w:ascii"), "Times New Roman")
            rFonts.set(qn("w:hAnsi"), "Times New Roman")
            rPr.append(rFonts)

            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), "24")
            rPr.append(sz)

            new_run.append(rPr)
            t_elem = OxmlElement("w:t")
            t_elem.text = "View Feedback Form"
            new_run.append(t_elem)
            hyperlink.append(new_run)
            link_para._p.append(hyperlink)

            paras_to_insert.append(link_para)

        # --- Warning if sheet data could not be fetched ---
        if fetch_failed:
            warn_para = doc.add_paragraph()
            warn_para.paragraph_format.space_before = Pt(10)
            warn_para.paragraph_format.space_after = Pt(8)
            warn_run = warn_para.add_run(
                "⚠ Could not fetch feedback data from Google Sheet. "
                "Please ensure the Google Sheet is shared publicly "
                "(File → Share → Publish to Web, or Share → Anyone with the link → Viewer) "
                "and try generating the report again."
            )
            _apply_font(warn_run, bold=True, size=Pt(11))
            from docx.shared import RGBColor
            warn_run.font.color.rgb = RGBColor(0xE1, 0x57, 0x59)
            paras_to_insert.append(warn_para)

        # --- 2. Summary statistics ---
        if stats.get("total", 0) > 0:
            summary_para = doc.add_paragraph()
            summary_para.paragraph_format.space_before = Pt(10)
            summary_para.paragraph_format.space_after = Pt(6)

            total_run = summary_para.add_run(
                f"Total Responses Received: {stats['total']}"
            )
            _apply_font(total_run, bold=True, size=Pt(12))

            if stats.get("charts_generated", 0) > 0:
                summary_para.add_run("\n")
                info_run = summary_para.add_run(
                    f"Analysis generated for {stats['charts_generated']} question(s)."
                )
                _apply_font(info_run, bold=False, size=Pt(11))

            paras_to_insert.append(summary_para)

        # --- 3. For each chart: sub-heading + image ---
        for ci in chart_info:
            chart_path = ci["path"]
            chart_title = ci["title"]

            if not os.path.exists(chart_path):
                continue

            # Sub-heading with the question text
            heading_para = doc.add_paragraph()
            heading_para.paragraph_format.space_before = Pt(14)
            heading_para.paragraph_format.space_after = Pt(4)
            heading_run = heading_para.add_run(chart_title)
            _apply_font(heading_run, bold=True, size=Pt(12))
            paras_to_insert.append(heading_para)

            # Chart image — centered, proper size
            img_para = doc.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_para.paragraph_format.keep_together = True
            img_para.paragraph_format.space_before = Pt(4)
            img_para.paragraph_format.space_after = Pt(14)
            run = img_para.add_run()
            run.add_picture(chart_path, width=Inches(5.2))
            paras_to_insert.append(img_para)

        # Move all created paragraphs from end-of-doc to after the anchor
        for para in reversed(paras_to_insert):
            anchor.addnext(para._p)

        # Remove the original marker paragraph
        anchor.getparent().remove(anchor)
        return


# ---------------- ROUTES ----------------

@app.route("/")
def home():
    return redirect(url_for("login"))


@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        username = request.form["username"]
        email = request.form.get("email", "")
        password = request.form["password"]
        role = request.form.get("role", "Student")  # Get selected role from form

        conn = get_db()
        existing = conn.execute(
            "SELECT id FROM users WHERE username=?", (username,)
        ).fetchone()
        if existing:
            conn.close()
            flash("Username already exists")
            return render_template("register.html")

        hashed = generate_password_hash(password)
        conn.execute(
            "INSERT INTO users (username, email, password, role) VALUES (?, ?, ?, ?)",
            (username, email, hashed, role),
        )
        conn.commit()
        conn.close()
        flash(f"Registration successful as {role}. Please login.")
        return redirect(url_for("login"))

    return render_template("register.html")


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form["username"]
        password = request.form["password"]
        selected_role = request.form.get("role", "Student")  # Get selected role from form

        conn = get_db()
        user = conn.execute(
            "SELECT * FROM users WHERE username=?", (username,)
        ).fetchone()
        conn.close()

        if user and check_password_hash(user["password"], password):
            # Get user's actual role from database
            user_role = user["role"] if "role" in user.keys() and user["role"] else "Student"
            
            # Verify that the selected role matches the user's role in database
            if selected_role != user_role:
                flash(f"Invalid role selection. Your account is registered as '{user_role}'. Please select the correct role.")
                return render_template("login.html")
            
            session["user_id"] = user["id"]
            session["username"] = user["username"]
            session["role"] = user_role
            return redirect(url_for("events"))
        flash("Invalid username or password")

    return render_template("login.html")


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))


@app.route("/events")
@login_required
def events():
    conn = get_db()
    rows = conn.execute("SELECT * FROM events ORDER BY id DESC").fetchall()
    conn.close()

    events = []
    for e in rows:
        event = dict(e)   # convert sqlite row to dict
        # Parse event_photos JSON and ensure it's a list
        try:
            event["photos"] = json.loads(e["event_photos"]) if e["event_photos"] else []
        except (json.JSONDecodeError, TypeError):
            event["photos"] = []
        # Ensure photos is always a list
        if not isinstance(event["photos"], list):
            event["photos"] = []
        events.append(event)

    return render_template("events.html", events=events)



@app.route("/add_event", methods=["GET", "POST"])
@login_required
def add_event():
    # Check if this is an edit request (event_id in query params or form)
    event_id = request.args.get("event_id", type=int) or request.form.get("event_id", type=int)
    event = None
    
    # Load existing event data if editing
    if event_id:
        conn = get_db()
        event = conn.execute("SELECT * FROM events WHERE id=?", (event_id,)).fetchone()
        conn.close()
        if not event:
            flash("Event not found.")
            return redirect(url_for("events"))
        event = dict(event)  # Convert to dict for easier access
    
    if request.method == "POST":
        title = request.form["title"]
        date = request.form["date"]
        venue = request.form["venue"]
        department = request.form["department"]
        description = request.form.get("description", "")
        academic_year = request.form.get("academic_year", "")
        resource_person = request.form.get("resource_person", "")
        resource_designation = request.form.get("resource_designation", "")
        event_coordinator = request.form.get("event_coordinator", "")
        event_time = request.form.get("event_time", "")
        event_type = request.form.get("event_type", "")
        outcome_1 = request.form.get("outcome_1", "")
        outcome_2 = request.form.get("outcome_2", "")
        outcome_3 = request.form.get("outcome_3", "")

        # PSO selections from form (checkboxes). Store as integers 0/1 in DB.
        pso1_selected = 1 if request.form.get("pso1_selected") else 0
        pso2_selected = 1 if request.form.get("pso2_selected") else 0

        # Validate: At least one PSO must be selected.
        if not (pso1_selected or pso2_selected):
            flash("Please select at least one Programme Specific Outcome (PSO1 and/or PSO2).")
            # Redirect back to the same form (preserves event_id for edit flows)
            if event_id:
                return redirect(url_for("add_event", event_id=event_id))
            return redirect(url_for("add_event"))

        # Program Outcomes (PO): multi-select; only allow hardcoded headings (no custom).
        raw_pos = request.form.getlist("selected_pos")
        selected_pos = [p for p in raw_pos if p in PO_HEADINGS]
        if not selected_pos:
            flash("Please select at least one Program Outcome.")
            if event_id:
                return redirect(url_for("add_event", event_id=event_id))
            return redirect(url_for("add_event"))

        # Sustainable Development Goals (SDGs): multi-select; optional (none = "Not Applicable" in report).
        valid_sdg_codes = {code for code, _ in SDG_LIST}
        raw_sdgs = request.form.getlist("selected_sdgs")
        selected_sdgs = [s for s in raw_sdgs if s in valid_sdg_codes]

        # Event photos (multiple) - preserve existing if no new uploads
        event_photos_paths = []
        if "event_photos" in request.files:
            for f in request.files.getlist("event_photos"):
                if f.filename:  # Only process if file was actually uploaded
                    saved = _save_file(f, "event_photos", allow_pdf=False)
                    if saved:
                        event_photos_paths.append(saved)
        
        # If editing and no new photos uploaded, keep existing photos
        if event_id and not event_photos_paths and event:
            try:
                existing_photos = json.loads(event["event_photos"]) if event.get("event_photos") else []
                if isinstance(existing_photos, list):
                    event_photos_paths = existing_photos
            except (json.JSONDecodeError, TypeError):
                pass

        event_photo_cover = event_photos_paths[0] if event_photos_paths else None

        # Scanned documents and attendance - preserve existing if no new uploads
        permission_letter = _save_file(
            request.files.get("permission_letter"), "attendance_photos", allow_pdf=True
        ) or (event["permission_letter"] if event_id and event else None)
        
        invitation_letter = _save_file(
            request.files.get("invitation_letter"), "attendance_photos", allow_pdf=True
        ) or (event["invitation_letter"] if event_id and event else None)
        
        notice_letter = _save_file(
            request.files.get("notice_letter"), "attendance_photos", allow_pdf=True
        ) or (event["notice_letter"] if event_id and event else None)
        
        appreciation_letter = _save_file(
            request.files.get("appreciation_letter"), "attendance_photos", allow_pdf=True
        ) or (event["appreciation_letter"] if event_id and event else None)
        
        attendance_photo = _save_file(
            request.files.get("attendance_photo"), "attendance_photos", allow_pdf=True
        ) or (event["attendance_photo"] if event_id and event else None)

        # Feedback — Google Form link + Google Sheet ID
        feedback_form_link = request.form.get("feedback_form_link", "").strip()
        feedback_sheet_id = request.form.get("feedback_sheet_id", "").strip()

        conn = get_db()
        
        # UPDATE if editing, INSERT if creating new
        if event_id:
            # Update existing event
            conn.execute(
                """
                UPDATE events SET
                    title=?, date=?, venue=?, department=?, description=?,
                    event_photo=?, academic_year=?, resource_person=?,
                    resource_designation=?, event_coordinator=?,
                    event_time=?, event_type=?,
                    permission_letter=?, invitation_letter=?, notice_letter=?,
                    appreciation_letter=?, event_photos=?, attendance_photo=?,
                    outcome_1=?, outcome_2=?, outcome_3=?,
                    pso1_selected=?, pso2_selected=?, selected_pos=?, selected_sdgs=?,
                    feedback_form_link=?, feedback_sheet_id=?
                WHERE id=?
                """,
                (
                    title,
                    date,
                    venue,
                    department,
                    description,
                    event_photo_cover,
                    academic_year,
                    resource_person,
                    resource_designation,
                    event_coordinator,
                    event_time,
                    event_type,
                    permission_letter,
                    invitation_letter,
                    notice_letter,
                    appreciation_letter,
                    json.dumps(event_photos_paths) if event_photos_paths else None,
                    attendance_photo,
                    outcome_1,
                    outcome_2,
                    outcome_3,
                    pso1_selected,
                    pso2_selected,
                    json.dumps(selected_pos),
                    json.dumps(selected_sdgs),
                    feedback_form_link or None,
                    feedback_sheet_id or None,
                    event_id,
                ),
            )
            flash("Event updated successfully.")
        else:
            # Insert new event
            conn.execute(
                """
                INSERT INTO events (
                    title, date, venue, department, description,
                    event_photo, academic_year, resource_person,
                    resource_designation, event_coordinator,
                    event_time, event_type,
                    permission_letter, invitation_letter, notice_letter,
                    appreciation_letter, event_photos, attendance_photo,
                    outcome_1, outcome_2, outcome_3,
                    pso1_selected, pso2_selected, selected_pos, selected_sdgs,
                    feedback_form_link, feedback_sheet_id
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    title,
                    date,
                    venue,
                    department,
                    description,
                    event_photo_cover,
                    academic_year,
                    resource_person,
                    resource_designation,
                    event_coordinator,
                    event_time,
                    event_type,
                    permission_letter,
                    invitation_letter,
                    notice_letter,
                    appreciation_letter,
                    json.dumps(event_photos_paths) if event_photos_paths else None,
                    attendance_photo,
                    outcome_1,
                    outcome_2,
                    outcome_3,
                    pso1_selected,
                    pso2_selected,
                    json.dumps(selected_pos),
                    json.dumps(selected_sdgs),
                    feedback_form_link or None,
                    feedback_sheet_id or None,
                ),
            )
            flash("Event added successfully.")
        
        conn.commit()
        conn.close()
        return redirect(url_for("events"))

    # GET request - render form with existing data if editing

    # Parse selected_pos for multi-select (edit mode)
    selected_pos_list = []
    if event and event.get("selected_pos"):
        try:
            selected_pos_list = json.loads(event["selected_pos"])
            if not isinstance(selected_pos_list, list):
                selected_pos_list = []
        except (json.JSONDecodeError, TypeError):
            selected_pos_list = []

    # Parse selected_sdgs for multi-select (edit mode)
    selected_sdgs_list = []
    if event and event.get("selected_sdgs"):
        try:
            selected_sdgs_list = json.loads(event["selected_sdgs"])
            if not isinstance(selected_sdgs_list, list):
                selected_sdgs_list = []
        except (json.JSONDecodeError, TypeError):
            selected_sdgs_list = []

    return render_template(
        "add_event.html",
        event=event,
        event_id=event_id,
        po_headings=PO_HEADINGS,
        selected_pos_list=selected_pos_list,
        sdg_list=SDG_LIST,
        selected_sdgs_list=selected_sdgs_list,
    )

@app.route("/event/<int:event_id>")
@login_required
def view_event(event_id):
    conn = get_db()
    event = conn.execute(
        "SELECT * FROM events WHERE id=?", (event_id,)
    ).fetchone()
    conn.close()

    if not event:
        flash("Event not found")
        return redirect(url_for("events"))

    photos = json.loads(event["event_photos"]) if event["event_photos"] else []

    return render_template("view_event.html", event=dict(event), photos=photos)


@app.route("/delete_event/<int:event_id>", methods=["POST"])
@login_required
def delete_event(event_id):
    """Delete an event from the database."""
    conn = get_db()
    event = conn.execute("SELECT * FROM events WHERE id=?", (event_id,)).fetchone()
    
    if not event:
        conn.close()
        flash("Event not found.")
        return redirect(url_for("events"))
    
    # Delete the event
    conn.execute("DELETE FROM events WHERE id=?", (event_id,))
    conn.commit()
    conn.close()
    
    flash("Event deleted successfully.")
    return redirect(url_for("events"))

@app.route("/generate_report/<int:event_id>")
@login_required
def generate_report(event_id):

    conn = get_db()
    event = conn.execute("SELECT * FROM events WHERE id=?", (event_id,)).fetchone()
    conn.close()

    if not event:
        flash("Event not found")
        return redirect(url_for("events"))

    # Convert sqlite3.Row to plain dict so we can safely use .get()
    event = dict(event)

    create_default_template()
    doc = Document(TEMPLATE_PATH)

    # Build PSO section text based on selected PSOs for this event.
    # This will replace the {{PSO_SECTION}} placeholder in the template.
    selected_psos = []
    if event.get("pso1_selected"):
        selected_psos.append(PSO1_TEXT)
    if event.get("pso2_selected"):
        selected_psos.append(PSO2_TEXT)
    pso_section_text = "\n\n".join(selected_psos)

    # Build PO section: selected Program Outcomes as bullet list (for {{PO_SECTION}}).
    # Equivalent to: {% for po in selected_pos %} • {{po}} {% endfor %}
    selected_pos = []
    if event.get("selected_pos"):
        try:
            selected_pos = json.loads(event["selected_pos"])
            if not isinstance(selected_pos, list):
                selected_pos = []
        except (json.JSONDecodeError, TypeError):
            selected_pos = []
    po_section_text = "\n".join("• " + po for po in selected_pos)

    # Build SDG section: selected SDGs as "SDG1: No Poverty" lines, or "Not Applicable".
    sdg_dict = dict(SDG_LIST)
    selected_sdgs = []
    if event.get("selected_sdgs"):
        try:
            selected_sdgs = json.loads(event["selected_sdgs"])
            if not isinstance(selected_sdgs, list):
                selected_sdgs = []
        except (json.JSONDecodeError, TypeError):
            selected_sdgs = []
    if not selected_sdgs:
        sdg_section_text = "Not Applicable"
    else:
        # Sort by SDG number so output is SDG1, SDG2, ...
        selected_sdgs_sorted = sorted(
            selected_sdgs,
            key=lambda c: int(c.replace("SDG", "")) if c.startswith("SDG") else 0,
        )
        sdg_section_text = "\n".join(
            f"{code}: {sdg_dict.get(code, code)}" for code in selected_sdgs_sorted
        )

    replace_placeholders(doc, {
        "{{academic_year}}": event["academic_year"],
        "{{date}}": event["date"],
        "{{event_name}}": event["title"],
        "{{event_type}}": event["event_type"],
        "{{event_date}}": event["date"],
        "{{event_time}}": event["event_time"],
        "{{venue}}": event["venue"],
        "{{department}}": event["department"],
        "{{resource_person}}": event["resource_person"],
        "{{resource_designation}}": event["resource_designation"],
        "{{event_coordinator}}": event["event_coordinator"],
        "{{event coordinator}}": event["event_coordinator"],
        # support both {{event_description}} and {{event description}} in template
        "{{event_description}}": event["description"],
        "{{event description}}": event["description"],
        "{{outcome_1}}": event["outcome_1"],
        "{{outcome_2}}": event["outcome_2"],
        "{{outcome_3}}": event["outcome_3"],
        "{{PSO_SECTION}}": pso_section_text,
        "{{PO_SECTION}}": po_section_text,
        "{{SDG_SECTION}}": sdg_section_text,
    })

    # --- Page break before Programme Outcome (POs) section ---
    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith("Programme Outcome"):
            run = OxmlElement("w:r")
            br = OxmlElement("w:br")
            br.set(qn("w:type"), "page")
            run.append(br)
            p._p.insert(0, run)
            break

    insert_event_details_paragraph(doc, "<<EVENT_DETAILS>>", event)

    # --- Each scanned document on its own separate page ---
    insert_full_page_image(doc, "<<IMAGE_PERMISSION>>", event["permission_letter"], heading="Permission Letter")
    insert_full_page_image(doc, "<<IMAGE_INVITATION>>", event["invitation_letter"], heading="Invitation Letter")
    insert_full_page_image(doc, "<<IMAGE_NOTICE>>", event["notice_letter"], heading="Notice")
    insert_full_page_image(doc, "<<IMAGE_APPRECIATION>>", event["appreciation_letter"], heading="Appreciation Letter")

    # --- Page break so main report content starts on a new page ---
    for p in doc.paragraphs:
        if p.text.strip().startswith("Department of"):
            run = OxmlElement("w:r")
            br = OxmlElement("w:br")
            br.set(qn("w:type"), "page")
            run.append(br)
            p._p.insert(0, run)
            break

    # --- Page break before Photographs section ---
    for p in doc.paragraphs:
        if p.text.strip().startswith("Photographs"):
            run = OxmlElement("w:r")
            br = OxmlElement("w:br")
            br.set(qn("w:type"), "page")
            run.append(br)
            p._p.insert(0, run)
            break

    photos = json.loads(event["event_photos"]) if event["event_photos"] else []
    insert_event_photos(doc, "<<EVENT_PHOTOS>>", photos)

    insert_attendance(doc, "<<ATTENDANCE_FILE>>", event["attendance_photo"])

    # --- Page break before Feedback section ---
    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith("Feedback") or text.startswith("Feedback Form"):
            run = OxmlElement("w:r")
            br = OxmlElement("w:br")
            br.set(qn("w:type"), "page")
            run.append(br)
            p._p.insert(0, run)
            break

    # Feedback Form & Analysis — fetch from Google Sheet + generate charts
    charts_dir = os.path.join(GENERATED_FOLDER, f"charts_{event_id}")
    insert_feedback_analysis(
        doc,
        "<<FEEDBACK_TABLE>>",
        event.get("feedback_form_link", ""),
        event.get("feedback_sheet_id", ""),
        charts_dir,
    )

    filename = event["title"].replace(" ", "_") + ".docx"
    path = os.path.join(GENERATED_FOLDER, filename)
    doc.save(path)

    # Store report metadata
    conn = get_db()
    conn.execute(
        # status column has a default ('submitted'), but be explicit for clarity.
        "INSERT INTO reports (event_id, file_path, status) VALUES (?, ?, ?)",
        (event_id, path, "submitted"),
    )
    conn.commit()
    conn.close()

    return send_file(path, as_attachment=True)


@app.route("/reports")
@login_required
def reports():
    conn = get_db()
    rows = conn.execute(
        """
        SELECT r.id,
               r.file_path,
               r.created_at,
               e.title AS event_title,
               e.date  AS event_date
        FROM reports r
        JOIN events e ON r.event_id = e.id
        ORDER BY r.created_at DESC
        """
    ).fetchall()
    conn.close()
    # Shape data to match template expectations
    reports_data = []
    for r in rows:
        reports_data.append(
            {
                "id": r["id"],
                "file_path": r["file_path"],
                "created_at": r["created_at"],
                "event_title": r["event_title"],
                "event_date": r["event_date"],
                "created_by_name": "",  # no separate user table linkage here
            }
        )
    return render_template("reports.html", reports=reports_data)


@app.route("/download_report/<int:report_id>")
@login_required
def download_report(report_id):
    conn = get_db()
    row = conn.execute(
        "SELECT file_path FROM reports WHERE id=?", (report_id,)
    ).fetchone()
    conn.close()
    if not row or not row["file_path"] or not os.path.exists(row["file_path"]):
        flash("Report file not found.")
        return redirect(url_for("reports"))
    return send_file(row["file_path"], as_attachment=True)


@app.route("/api/hod/department-analysis")
@login_required
@hod_required_api
def hod_department_analysis_api():
    """
    Returns aggregated, department-wise report analytics for HOD users.

    Output format (example):
    {
        "departments": [
            {
                "department": "CSE",
                "total_reports": 10,
                "status_counts": {
                    "submitted": 5,
                    "approved": 3,
                    "rejected": 1,
                    "pending": 1
                }
            },
            ...
        ]
    }
    """
    conn = get_db()

    # Aggregate per department and per status. We keep the schema flexible while
    # still returning a shape that is easy for charting libraries to consume.
    rows = conn.execute(
        """
        SELECT
            e.department AS department,
            r.status     AS status,
            COUNT(r.id)  AS count
        FROM reports r
        JOIN events e ON r.event_id = e.id
        GROUP BY e.department, r.status
        ORDER BY e.department
        """
    ).fetchall()
    conn.close()

    # Build nested structure: department -> status_counts + total
    departments = {}
    for row in rows:
        dept = row["department"] or "Unknown"
        status = row["status"] or "unknown"
        count = row["count"] or 0

        if dept not in departments:
            departments[dept] = {
                "department": dept,
                "total_reports": 0,
                "status_counts": {},
            }
        departments[dept]["status_counts"][status] = (
            departments[dept]["status_counts"].get(status, 0) + count
        )
        departments[dept]["total_reports"] += count

    # Sorted list is a bit nicer for charts/UX
    department_list = sorted(
        departments.values(), key=lambda d: d["department"].lower()
    )

    return jsonify({"departments": department_list})


# ---------------- START ----------------

import threading
import webbrowser

def open_browser():
    webbrowser.open("http://127.0.0.1:5000")

if __name__ == "__main__":
    threading.Timer(1, open_browser).start()
    app.run(debug=False)


